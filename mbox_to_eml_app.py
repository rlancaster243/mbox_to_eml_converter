import streamlit as st
import tempfile
import os
import io
import pandas as pd
from PIL import Image
import markdown as md
import nbformat
import mailbox
import zipfile

try:
    from pdf2docx import Converter
    import pdfplumber
    from pdf2image import convert_from_bytes
    import pytesseract
except ImportError:
    pass

# Supported conversions (Streamlit Cloud compatible only)
CONVERSION_MATRIX = {
    "txt": ["docx", "md", "html"],
    "md": ["html", "txt"],
    "html": ["txt", "md"],
    "csv": ["xlsx", "txt"],
    "tsv": ["xlsx", "txt"],
    "xlsx": ["csv", "tsv", "txt"],
    "docx": ["txt"],
    "pdf": ["docx", "xlsx", "txt"],
    "jpg": ["png", "bmp", "gif", "tiff", "webp"],
    "png": ["jpg", "bmp", "gif", "tiff", "webp"],
    "bmp": ["jpg", "png", "gif", "tiff", "webp"],
    "gif": ["jpg", "png", "bmp", "tiff", "webp"],
    "tiff": ["jpg", "png", "bmp", "gif", "webp"],
    "webp": ["jpg", "png", "bmp", "gif", "tiff"],
    "mbox": ["eml"],
    "eml": ["mbox"],
    "py": ["ipynb"],
    "ipynb": ["py"],
}

EXTENSION_LABELS = {
    "txt": "Plain Text",
    "md": "Markdown",
    "html": "HTML",
    "csv": "CSV",
    "tsv": "TSV",
    "xlsx": "Excel (XLSX)",
    "docx": "Word (DOCX)",
    "pdf": "PDF",
    "jpg": "JPEG Image",
    "png": "PNG Image",
    "bmp": "Bitmap Image",
    "gif": "GIF Image",
    "tiff": "TIFF Image",
    "webp": "WEBP Image",
    "mbox": "MBOX Email Mailbox",
    "eml": "EML Email Message",
    "py": "Python Script",
    "ipynb": "Jupyter Notebook",
}

# -------- Conversion Functions --------

def txt_to_docx(txt_bytes):
    from docx import Document
    text = txt_bytes.decode('utf-8')
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    output = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(output.name)
    with open(output.name, "rb") as f:
        result = f.read()
    os.remove(output.name)
    return result

def txt_to_md(txt_bytes):
    return txt_bytes

def txt_to_html(txt_bytes):
    text = txt_bytes.decode('utf-8')
    html = "<br>".join(text.splitlines())
    return html.encode('utf-8')

def md_to_html(md_bytes):
    text = md_bytes.decode('utf-8')
    html = md.markdown(text)
    return html.encode('utf-8')

def md_to_txt(md_bytes):
    return md_bytes

def html_to_txt(html_bytes):
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_bytes.decode('utf-8'), "html.parser")
    return soup.get_text().encode('utf-8')

def html_to_md(html_bytes):
    from markdownify import markdownify as mdify
    mdtext = mdify(html_bytes.decode('utf-8'))
    return mdtext.encode('utf-8')

def csv_to_xlsx(csv_bytes):
    df = pd.read_csv(io.BytesIO(csv_bytes))
    output = io.BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)
    return output.read()

def csv_to_txt(csv_bytes):
    return csv_bytes

def tsv_to_xlsx(tsv_bytes):
    df = pd.read_csv(io.BytesIO(tsv_bytes), sep='\t')
    output = io.BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)
    return output.read()

def tsv_to_txt(tsv_bytes):
    return tsv_bytes

def xlsx_to_csv(xlsx_bytes):
    df = pd.read_excel(io.BytesIO(xlsx_bytes))
    return df.to_csv(index=False).encode('utf-8')

def xlsx_to_tsv(xlsx_bytes):
    df = pd.read_excel(io.BytesIO(xlsx_bytes))
    return df.to_csv(index=False, sep='\t').encode('utf-8')

def xlsx_to_txt(xlsx_bytes):
    df = pd.read_excel(io.BytesIO(xlsx_bytes))
    return df.to_string(index=False).encode('utf-8')

def docx_to_txt(docx_bytes):
    from docx import Document
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tf:
        tf.write(docx_bytes)
        tf.flush()
        doc = Document(tf.name)
    text = "\n".join([p.text for p in doc.paragraphs])
    os.remove(tf.name)
    return text.encode('utf-8')

def pdf_to_docx(pdf_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tf:
        tf.write(pdf_bytes)
        tf.flush()
        docx_path = tf.name.replace('.pdf', '.docx')
        cv = Converter(tf.name)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
    with open(docx_path, 'rb') as f:
        output_bytes = f.read()
    os.remove(docx_path)
    return output_bytes

def pdf_to_xlsx(pdf_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tf:
        tf.write(pdf_bytes)
        tf.flush()
        with pdfplumber.open(tf.name) as pdf:
            all_tables = []
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    all_tables.append(table)
        if not all_tables:
            return None
        writer = pd.ExcelWriter(tf.name.replace('.pdf', '.xlsx'))
        for i, table in enumerate(all_tables):
            df = pd.DataFrame(table[1:], columns=table[0])
            df.to_excel(writer, index=False, sheet_name=f"Table_{i+1}")
        writer.close()
        with open(tf.name.replace('.pdf', '.xlsx'), 'rb') as f:
            xlsx_bytes = f.read()
        os.remove(tf.name.replace('.pdf', '.xlsx'))
    return xlsx_bytes

def pdf_to_text_ocr(pdf_bytes):
    images = convert_from_bytes(pdf_bytes)
    text = ""
    for i, img in enumerate(images):
        text += f"\n--- PAGE {i+1} ---\n"
        text += pytesseract.image_to_string(img)
    return text.encode('utf-8')

def image_convert(image_bytes, to_ext):
    with Image.open(io.BytesIO(image_bytes)) as img:
        output = io.BytesIO()
        img.save(output, format=to_ext.upper())
        output.seek(0)
        return output.read()

def mbox_to_eml_files(mbox_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.mbox') as tf:
        tf.write(mbox_bytes)
        tf.flush()
        mbox = mailbox.mbox(tf.name)
        eml_files = []
        for i, msg in enumerate(mbox, 1):
            eml_bytes = bytes(msg)
            eml_files.append((f"{i:04d}.eml", eml_bytes))
    os.remove(tf.name)
    return eml_files

def eml_files_to_mbox(eml_file_list):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.mbox') as tf:
        mbox = mailbox.mbox(tf.name)
        for name, eml_bytes in eml_file_list:
            msg = mailbox.mboxMessage(eml_bytes)
            mbox.add(msg)
        mbox.flush()
        with open(tf.name, 'rb') as mbox_file:
            mbox_bytes = mbox_file.read()
    os.remove(tf.name)
    return mbox_bytes

def py_to_ipynb(py_bytes):
    code = py_bytes.decode('utf-8')
    notebook = nbformat.v4.new_notebook()
    notebook.cells.append(nbformat.v4.new_code_cell(code))
    return nbformat.writes(notebook).encode('utf-8')

def ipynb_to_py(ipynb_bytes):
    notebook = nbformat.reads(ipynb_bytes.decode('utf-8'), as_version=4)
    code = ""
    for cell in notebook.cells:
        if cell.cell_type == 'code':
            code += cell.source + '\n\n'
    return code.encode('utf-8')

# ---- Streamlit UI ----

st.title("Universal File Converter (Streamlit Cloud Compatible)")

from_ext = st.selectbox(
    "Convert FROM",
    options=list(CONVERSION_MATRIX.keys()),
    format_func=lambda x: EXTENSION_LABELS.get(x, x)
)
to_ext_options = CONVERSION_MATRIX[from_ext]
to_ext = st.selectbox(
    "Convert TO",
    options=to_ext_options,
    format_func=lambda x: EXTENSION_LABELS.get(x, x) if x in EXTENSION_LABELS else x
)

accept_multiple = from_ext in [
    "mbox", "eml", "jpg", "png", "bmp", "gif", "tiff", "webp"
]
uploaded_files = st.file_uploader(
    f"Upload .{from_ext} file(s)",
    type=[from_ext],
    accept_multiple_files=accept_multiple
)

st.info(f"Selected: {EXTENSION_LABELS.get(from_ext, from_ext)} → {EXTENSION_LABELS.get(to_ext, to_ext)}")

if uploaded_files and st.button("Convert & Download"):
    try:
        # Text/Markdown/HTML
        if from_ext == "txt" and to_ext == "docx":
            output_bytes = txt_to_docx(uploaded_files[0].read())
            st.download_button("Download DOCX", data=output_bytes, file_name="converted.docx")
        elif from_ext == "txt" and to_ext == "md":
            st.download_button("Download MD", data=txt_to_md(uploaded_files[0].read()), file_name="converted.md")
        elif from_ext == "txt" and to_ext == "html":
            st.download_button("Download HTML", data=txt_to_html(uploaded_files[0].read()), file_name="converted.html")
        elif from_ext == "md" and to_ext == "html":
            st.download_button("Download HTML", data=md_to_html(uploaded_files[0].read()), file_name="converted.html")
        elif from_ext == "md" and to_ext == "txt":
            st.download_button("Download TXT", data=md_to_txt(uploaded_files[0].read()), file_name="converted.txt")
        elif from_ext == "html" and to_ext == "txt":
            st.download_button("Download TXT", data=html_to_txt(uploaded_files[0].read()), file_name="converted.txt")
        elif from_ext == "html" and to_ext == "md":
            st.download_button("Download MD", data=html_to_md(uploaded_files[0].read()), file_name="converted.md")

        # Spreadsheet
        elif from_ext == "csv" and to_ext == "xlsx":
            st.download_button("Download XLSX", data=csv_to_xlsx(uploaded_files[0].read()), file_name="converted.xlsx")
        elif from_ext == "csv" and to_ext == "txt":
            st.download_button("Download TXT", data=csv_to_txt(uploaded_files[0].read()), file_name="converted.txt")
        elif from_ext == "tsv" and to_ext == "xlsx":
            st.download_button("Download XLSX", data=tsv_to_xlsx(uploaded_files[0].read()), file_name="converted.xlsx")
        elif from_ext == "tsv" and to_ext == "txt":
            st.download_button("Download TXT", data=tsv_to_txt(uploaded_files[0].read()), file_name="converted.txt")
        elif from_ext == "xlsx" and to_ext == "csv":
            st.download_button("Download CSV", data=xlsx_to_csv(uploaded_files[0].read()), file_name="converted.csv")
        elif from_ext == "xlsx" and to_ext == "tsv":
            st.download_button("Download TSV", data=xlsx_to_tsv(uploaded_files[0].read()), file_name="converted.tsv")
        elif from_ext == "xlsx" and to_ext == "txt":
            st.download_button("Download TXT", data=xlsx_to_txt(uploaded_files[0].read()), file_name="converted.txt")

        # DOCX → TXT
        elif from_ext == "docx" and to_ext == "txt":
            st.download_button("Download TXT", data=docx_to_txt(uploaded_files[0].read()), file_name="converted.txt")

        # PDF (all three)
        elif from_ext == "pdf" and to_ext == "docx":
            st.download_button("Download DOCX", data=pdf_to_docx(uploaded_files[0].read()), file_name="converted.docx")
        elif from_ext == "pdf" and to_ext == "xlsx":
            st.download_button("Download XLSX", data=pdf_to_xlsx(uploaded_files[0].read()), file_name="converted.xlsx")
        elif from_ext == "pdf" and to_ext == "txt":
            st.download_button("Download OCR TXT", data=pdf_to_text_ocr(uploaded_files[0].read()), file_name="ocr_output.txt")

        # Images
        elif from_ext in ["jpg", "png", "bmp", "gif", "tiff", "webp"] and to_ext in ["jpg", "png", "bmp", "gif", "tiff", "webp"]:
            outputs = []
            for f in uploaded_files:
                converted = image_convert(f.read(), to_ext)
                outputs.append((os.path.splitext(f.name)[0] + "." + to_ext, converted))
            if len(outputs) == 1:
                st.download_button(f"Download {to_ext.upper()}", data=outputs[0][1], file_name=outputs[0][0])
            else:
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zipf:
                    for fname, content in outputs:
                        zipf.writestr(fname, content)
                zip_buffer.seek(0)
                st.download_button(f"Download All as ZIP", data=zip_buffer, file_name=f"converted_{to_ext}s.zip")

        # mbox <-> eml
        elif from_ext == "mbox" and to_ext == "eml":
            eml_files = mbox_to_eml_files(uploaded_files[0].read())
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                for fname, eml_bytes in eml_files:
                    zipf.writestr(fname, eml_bytes)
            zip_buffer.seek(0)
            st.download_button("Download EML Zip", data=zip_buffer, file_name="emls.zip")
        elif from_ext == "eml" and to_ext == "mbox":
            eml_file_list = [(f.name, f.read()) for f in uploaded_files]
            mbox_bytes = eml_files_to_mbox(eml_file_list)
            st.download_button("Download MBOX", data=mbox_bytes, file_name="converted.mbox")

        # py <-> ipynb
        elif from_ext == "py" and to_ext == "ipynb":
            st.download_button("Download Notebook", data=py_to_ipynb(uploaded_files[0].read()), file_name="converted.ipynb")
        elif from_ext == "ipynb" and to_ext == "py":
            st.download_button("Download Script", data=ipynb_to_py(uploaded_files[0].read()), file_name="converted.py")

        else:
            st.error("This conversion is not supported on Streamlit Cloud (requires system dependencies).")
    except Exception as e:
        st.error(f"Error during conversion: {e}")

with st.expander("Supported File Types and Conversion Matrix"):
    st.write(CONVERSION_MATRIX)
